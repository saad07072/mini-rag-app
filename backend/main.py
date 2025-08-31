import os
from fastapi import FastAPI, UploadFile, HTTPException, status
from pydantic import BaseModel
from typing import Union, List
import uuid
from sentence_transformers import SentenceTransformer, CrossEncoder
from qdrant_client import QdrantClient, models
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
from io import BytesIO
from dotenv import load_dotenv
from groq import Groq

# Load environment variables from a .env file
load_dotenv()

app = FastAPI()

# ---- Groq Client ----
# Get API key from environment variables or hardcode for dev purposes
groq_client = Groq(api_key=os.getenv("GROQ_API_KEY", "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"))

# ---- Qdrant Client ----
qdrant_client = QdrantClient(
    url=os.getenv("https://7794553c-496c-4037-91b1-b0c4a0f56f25.us-west-1-0.aws.cloud.qdrant.io"),
    api_key=os.getenv("eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJhY2Nlc3MiOiJtIn0.8w7_q0dD5tEQnzl45YjcUKJSj4WRst9lYQDGlN-Fxng")
)
collection_name = "documents"

# ---- Sentence Transformer Models ----
# This model is loaded once at app startup for efficiency.
embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
# Reranker model for improving search result relevance
reranker_model = CrossEncoder("cross-encoder/ms-marco-MiniLM-L-6-v2")


def get_embedding(text: str):
    """Generates an embedding vector for the given text."""
    return embedding_model.encode(text).tolist()

def extract_text_from_file(file: UploadFile) -> str:
    """Extracts text from different file types."""
    content = ""
    file_type = file.filename.split('.')[-1].lower()
    
    if file_type == 'txt':
        content = file.file.read().decode('utf-8')
    elif file_type == 'pdf':
        reader = PdfReader(BytesIO(file.file.read()))
        for page in reader.pages:
            content += page.extract_text() or ""
    elif file_type == 'docx':
        doc = DocxDocument(BytesIO(file.file.read()))
        for para in doc.paragraphs:
            content += para.text + "\n"
    else:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsupported file type")
    
    return content

# ---- Pydantic Models ----
class Document(BaseModel):
    text: str
    id: Union[int, str, None] = None

class Query(BaseModel):
    text: str
    top_k: int = 3
    
class SearchResult(BaseModel):
    id: str
    text: str
    score: float

class LLMResponse(BaseModel):
    answer: str
    sources: List[str]

# ---- Endpoints ----
@app.post("/add_document")
async def add_document(doc: Document):
    vector = get_embedding(doc.text)

    # Auto-generate ID if not provided
    if doc.id is None:
        point_id = str(uuid.uuid4())
    else:
        point_id = int(doc.id) if str(doc.id).isdigit() else str(uuid.UUID(str(doc.id)))

    qdrant_client.upsert(
        collection_name=collection_name,
        points=[
            models.PointStruct(
                id=point_id,
                vector=vector,
                payload={"text": doc.text}
            )
        ]
    )
    return {"status": "success", "id": point_id}

@app.post("/upload_document")
async def upload_document(file: UploadFile):
    content = extract_text_from_file(file)
    vector = get_embedding(content)
    point_id = str(uuid.uuid4())
    
    qdrant_client.upsert(
        collection_name=collection_name,
        points=[
            models.PointStruct(
                id=point_id,
                vector=vector,
                payload={"text": content, "filename": file.filename}
            )
        ]
    )
    return {"filename": file.filename, "id": point_id, "status": "success"}

@app.post("/query")
async def query_documents(query: Query) -> List[SearchResult]:
    query_vector = get_embedding(query.text)
    
    search_result = qdrant_client.search(
        collection_name=collection_name,
        query_vector=query_vector,
        limit=query.top_k
    )

    results = []
    for point in search_result:
        results.append(SearchResult(
            id=str(point.id),
            text=point.payload['text'],
            score=point.score
        ))
    return results

@app.post("/generate_answer")
async def generate_answer(query: Query) -> LLMResponse:
    # Step 1: Retrieve relevant documents
    retrieved_results = qdrant_client.search(
        collection_name=collection_name,
        query_vector=get_embedding(query.text),
        limit=query.top_k * 2 # Retrieve more documents for better reranking
    )
    
    # Step 2: Rerank the retrieved documents
    rerank_candidates = [point.payload['text'] for point in retrieved_results]
    reranked_scores = reranker_model.predict([(query.text, doc) for doc in rerank_candidates])
    
    # Sort documents by their new reranked score
    reranked_results = sorted(zip(rerank_candidates, reranked_scores), key=lambda x: x[1], reverse=True)
    
    # Step 3: Select the top 3 reranked documents for the LLM context
    top_reranked_docs = reranked_results[:query.top_k]

    # Step 4: Combine documents into a single context
    context = ""
    for i, (text, score) in enumerate(top_reranked_docs):
        context += f"Document {i+1}: {text}\n\n"

    # Step 5: Construct the prompt for the LLM
    prompt = f"""
    Based on the following documents, answer the user's question. If the information is not present in the documents, state that you cannot answer. Do not use any external knowledge.

    Documents:
    ---
    {context}
    ---

    User's Question: {query.text}

    Answer:
    """

    # Step 6: Call the Groq API to generate an answer
    chat_completion = groq_client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": prompt,
            }
        ],
        model="llama3-8b-8192",
        temperature=0.5,
        max_tokens=1024
    )
    
    llm_answer = chat_completion.choices[0].message.content
    
    # Step 7: Return the generated answer and the source documents
    sources = [text for text, score in top_reranked_docs]
    
    return LLMResponse(answer=llm_answer, sources=sources)